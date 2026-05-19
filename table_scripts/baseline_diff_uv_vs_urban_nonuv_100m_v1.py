#!/usr/bin/env python3
"""
Descriptive Statistics: Five-sample baseline differences (100m grid level)
Output: CSV and docx table with grid counts and period counts
"""

import os
import pandas as pd
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Setup paths
ROOT = r"e:\Codex\Tariff_shock_crime_and_Infrastructure"
OUT_TABLE = os.path.join(ROOT, "table/main")
os.makedirs(OUT_TABLE, exist_ok=True)

# Load data
data_csv = os.path.join(ROOT, "grid_halfyear_panel_100m_judicial_exposure_v3_with_housing_noradiusmerge_v3.csv")
print(f"Loading CSV: {data_csv}")
df = pd.read_csv(data_csv, encoding='utf-8')

print(f"Data dimensions: {df.shape[0]} rows, {df.shape[1]} cols")
print(f"Column names: {', '.join(df.columns[:10])}...")

# Ensure gridid exists
if 'gridid' not in df.columns:
    if 'cell_id' in df.columns:
        df['gridid'] = df['cell_id'].astype(str)
        print("Created gridid from cell_id")
    elif 'x' in df.columns and 'y' in df.columns:
        df['gridid'] = df['x'].astype(str) + "_" + df['y'].astype(str)
        print("Created gridid from x_y")
    else:
        raise ValueError("No gridid, cell_id, or x/y found in data")

# Create derived variables
df['year_num'] = df['period'].str[:4].astype(int)
df['half_num'] = df['period'].str[5].astype(int)
df['post'] = (df['year_num'] > 2018) | ((df['year_num'] == 2018) & (df['half_num'] >= 2))
df['is_urban'] = pd.to_numeric(df['is_urban'], errors='coerce')
df['is_uv'] = pd.to_numeric(df['is_uv'], errors='coerce')
df['metro5'] = df['county_city'].isin(['广州市', '深圳市', '佛山市', '东莞市', '珠海市'])

# Variable list
varlist = [
    "us_tariff_exposure_4", "pop", "crime_rate_100k", "pct_urban_village",
    "pre_avg_list_unit_price", "pre_avg_deal_unit_price", "pre_avg_price_gap",
    "dist_to_bus_m", "dist_to_metro_m",
    "dist_to_edu_pre_m", "dist_to_edu_primary_m", "dist_to_edu_secondary_m",
    "dist_to_edu_higher_m", "dist_to_edu_vocational_m",
    "dist_to_basic_healthcare_m", "dist_to_mid_healthcare_m", "dist_to_adv_healthcare_m",
    "dist_to_sme_bank_m", "dist_to_joint_stock_bank_m", "dist_to_state_owned_bank_m"
]

# Filter to only existing variables
varlist = [v for v in varlist if v in df.columns]
print(f"Processing {len(varlist)} variables")

# Process each variable
results = []

for v in varlist:
    print(f"Processing: {v}")
    
    # Define five samples (masks)
    mask_full = df[v].notna()
    mask_urban = (df['is_urban'] == 1) & (df['is_uv'] != 1) & (df[v].notna())
    mask_uv = (df['is_uv'] == 1) & (df[v].notna())
    mask_center = (df['is_urban'] == 1) & (df['is_uv'] == 1) & (df[v].notna())
    mask_periph = (df['is_urban'] == 0) & (df['is_uv'] == 1) & (df[v].notna())
    
    # Handle dist_to_metro_m: restrict to metro5 only
    if v == "dist_to_metro_m":
        mask_full = mask_full & df['metro5']
        mask_urban = mask_urban & df['metro5']
        mask_uv = mask_uv & df['metro5']
        mask_center = mask_center & df['metro5']
        mask_periph = mask_periph & df['metro5']
    
    # Compute means
    mean_full = df.loc[mask_full, v].mean()
    mean_urban = df.loc[mask_urban, v].mean()
    mean_uv = df.loc[mask_uv, v].mean()
    mean_center = df.loc[mask_center, v].mean()
    mean_periph = df.loc[mask_periph, v].mean()
    
    # Compute unique grid counts
    grids_full = df.loc[mask_full, 'gridid'].nunique()
    grids_urban = df.loc[mask_urban, 'gridid'].nunique()
    grids_uv = df.loc[mask_uv, 'gridid'].nunique()
    grids_center = df.loc[mask_center, 'gridid'].nunique()
    grids_periph = df.loc[mask_periph, 'gridid'].nunique()
    
    # Compute distinct periods
    periods_full = df.loc[mask_full, 'period'].nunique()
    periods_urban = df.loc[mask_urban, 'period'].nunique()
    periods_uv = df.loc[mask_uv, 'period'].nunique()
    periods_center = df.loc[mask_center, 'period'].nunique()
    periods_periph = df.loc[mask_periph, 'period'].nunique()
    
    results.append({
        'varname': v,
        'vlabel': v,
        'mean_full': mean_full,
        'mean_urban': mean_urban,
        'mean_uv': mean_uv,
        'mean_center': mean_center,
        'mean_periph': mean_periph,
        'grids_full': grids_full,
        'grids_urban': grids_urban,
        'grids_uv': grids_uv,
        'grids_center': grids_center,
        'grids_periph': grids_periph,
        'periods_full': periods_full,
        'periods_urban': periods_urban,
        'periods_uv': periods_uv,
        'periods_center': periods_center,
        'periods_periph': periods_periph,
    })

# Combine results
res_df = pd.DataFrame(results)

# Export CSV
csv_out = os.path.join(OUT_TABLE, "baseline_diff_uv_vs_urban_nonuv_100m_v1.csv")
res_df.to_csv(csv_out, index=False, encoding='utf-8')
print(f"Exported CSV: {csv_out}")

# Compute overall grid/period counts (using full dataset)
mask_full_all = df['gridid'].notna()
mask_urban_all = (df['is_urban'] == 1) & (df['is_uv'] != 1) & (df['gridid'].notna())
mask_uv_all = (df['is_uv'] == 1) & (df['gridid'].notna())
mask_center_all = (df['is_urban'] == 1) & (df['is_uv'] == 1) & (df['gridid'].notna())
mask_periph_all = (df['is_urban'] == 0) & (df['is_uv'] == 1) & (df['gridid'].notna())

grids_all = df.loc[mask_full_all, 'gridid'].nunique()
grids_urban_all = df.loc[mask_urban_all, 'gridid'].nunique()
grids_uv_all = df.loc[mask_uv_all, 'gridid'].nunique()
grids_center_all = df.loc[mask_center_all, 'gridid'].nunique()
grids_periph_all = df.loc[mask_periph_all, 'gridid'].nunique()

periods_all = df.loc[mask_full_all, 'period'].nunique()
periods_urban_all = df.loc[mask_urban_all, 'period'].nunique()
periods_uv_all = df.loc[mask_uv_all, 'period'].nunique()
periods_center_all = df.loc[mask_center_all, 'period'].nunique()
periods_periph_all = df.loc[mask_periph_all, 'period'].nunique()

print("Grid and period counts computed")

# Create docx table
doc = Document()

# Add title
title = doc.add_paragraph("Descriptive Statistics: Five-sample grid-level counts + means (100m)")
title.runs[0].bold = True
title.runs[0].font.size = Pt(14)

# Prepare table data for display
table_rows = []
for _, row in res_df.iterrows():
    table_rows.append([
        row['varname'],
        f"{row['mean_full']:.3f}" if pd.notna(row['mean_full']) else "",
        f"{row['mean_urban']:.3f}" if pd.notna(row['mean_urban']) else "",
        f"{row['mean_uv']:.3f}" if pd.notna(row['mean_uv']) else "",
        f"{row['mean_center']:.3f}" if pd.notna(row['mean_center']) else "",
        f"{row['mean_periph']:.3f}" if pd.notna(row['mean_periph']) else "",
    ])

# Create table with header
table = doc.add_table(rows=len(table_rows) + 1, cols=6)
table.style = 'Light Grid Accent 1'

# Add header
header_cells = table.rows[0].cells
header_texts = ["Variable", "全样本", "正规街区 (urban=1, uv=0)", "城中村 (uv=1)", 
                "中心城中村 (urban=1 & uv=1)", "外围城中村 (urban=0 & uv=1)"]
for i, text in enumerate(header_texts):
    header_cells[i].text = text

# Add data rows
for i, row_data in enumerate(table_rows, start=1):
    row_cells = table.rows[i].cells
    for j, text in enumerate(row_data):
        row_cells[j].text = str(text)

# Add observations row
obs_row = table.add_row().cells
obs_row[0].text = "Observations (unique grids)"
obs_row[1].text = str(grids_all)
obs_row[2].text = str(grids_urban_all)
obs_row[3].text = str(grids_uv_all)
obs_row[4].text = str(grids_center_all)
obs_row[5].text = str(grids_periph_all)

# Add periods row
periods_row = table.add_row().cells
periods_row[0].text = "Periods"
periods_row[1].text = str(periods_all)
periods_row[2].text = str(periods_urban_all)
periods_row[3].text = str(periods_uv_all)
periods_row[4].text = str(periods_center_all)
periods_row[5].text = str(periods_periph_all)

# Add notes
notes = doc.add_paragraph(
    "Notes: Means are observation-level means; Observations are unique grid counts per sample; "
    "Periods are distinct half-year periods in sample; dist_to_metro_m restricted to metro5 cities."
)
notes.paragraph_format.space_before = Pt(12)

# Save docx
docx_out = os.path.join(OUT_TABLE, "baseline_diff_uv_vs_urban_nonuv_100m_v1.docx")
doc.save(docx_out)
print(f"Exported DOCX: {docx_out}")

# Print summary
print("\n=== Summary ===")
print(f"Variables processed: {len(res_df)}")
print(f"Grid counts (full sample): {grids_all}")
print(f"Period count (full sample): {periods_all}")
print("Done: baseline difference table exported (csv + docx) with five columns and grid counts.")

# Print first few rows
print("\n=== First 5 rows of CSV ===")
print(res_df.head().to_string())
